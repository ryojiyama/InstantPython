#!/usr/bin/env python
"""
xlsx_to_docx.py  ―  DrawingML 解析による xlsx → docx 一括変換

使い方:
    python xlsx_to_docx.py

ディレクトリ構造:
    officedoc_transfer/
    ├── xlsx_to_docx.py   ← このファイル
    ├── input/            ← 変換したい .xlsx をここに置く（複数可）
    └── output/           ← 変換後の .docx がここに出力される
                             └── {元ファイル名（拡張子なし）}/
                                 └── {シート名}.docx

処理対象シートの条件:
    drawing XML にテキストまたは画像アンカーが 1 件以上あるシートのみ。
    空シート・drawing なしシートは自動スキップ。

依存ライブラリ:
    pip install python-docx
"""

import zipfile
import xml.etree.ElementTree as ET
import re
import io
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu

# ─── 名前空間 ──────────────────────────────────────────────────────────────────
NS_XDR  = 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing'
NS_A    = 'http://schemas.openxmlformats.org/drawingml/2006/main'
NS_R    = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
NS_PKG  = 'http://schemas.openxmlformats.org/package/2006/relationships'
NS_MAIN = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

def q(ns, local):
    return f'{{{ns}}}{local}'

# ─── ページ設定 ────────────────────────────────────────────────────────────────
A4_W          = Cm(21.0)
A4_H          = Cm(29.7)
MARGIN_TOP    = Cm(1.5)
MARGIN_BOTTOM = Cm(2.0)   # フッター領域を考慮して広め
MARGIN_LEFT   = Cm(1.8)
MARGIN_RIGHT  = Cm(1.8)
CONTENT_W_EMU = int(A4_W - MARGIN_LEFT - MARGIN_RIGHT)

# ─── 見出し検出パターン ────────────────────────────────────────────────────────
HEADING_PATTERN = re.compile(
    r'^[１２３４５６７８９\d][）\)]'   # 「１）」「1)」形式の工程番号
    r'|^【.+】'                          # 「【部品名】」形式
)

# ─── ユーティリティ ────────────────────────────────────────────────────────────
def sz_to_pt(sz_str):
    """DrawingML のフォントサイズ（1/100pt 単位）を pt に変換"""
    if sz_str is None:
        return 10.5
    return int(sz_str) / 100.0

def safe_filename(name):
    """OS で使えない文字をアンダースコアに置換"""
    return re.sub(r'[\\/:*?"<>|]', '_', name)

def log(msg, indent=0):
    print('  ' * indent + msg)

# ─── xlsx 解析 ─────────────────────────────────────────────────────────────────
def get_sheets(zf):
    """workbook.xml からシート情報（名前・シートファイルパス）のリストを返す"""
    with zf.open("xl/workbook.xml") as f:
        wb = ET.parse(f).getroot()
    with zf.open("xl/_rels/workbook.xml.rels") as f:
        rels_root = ET.parse(f).getroot()

    rid_to_file = {
        rel.get('Id'): rel.get('Target')
        for rel in rels_root.findall(q(NS_PKG, 'Relationship'))
        if 'worksheet' in rel.get('Type', '')
    }

    return [
        {
            'name': s.get('name'),
            'sheet_file': rid_to_file.get(s.get(q(NS_R, 'id')), ''),
        }
        for s in wb.findall(f'.//{q(NS_MAIN, "sheet")}')
    ]


def get_drawing_file(zf, sheet_file):
    """シートに対応する drawing XML のパスを返す。なければ None"""
    rels_path = 'xl/' + sheet_file.replace('worksheets/', 'worksheets/_rels/') + '.rels'
    if rels_path not in zf.namelist():
        return None
    with zf.open(rels_path) as f:
        rels_root = ET.parse(f).getroot()
    for rel in rels_root.findall(q(NS_PKG, 'Relationship')):
        if 'drawing' in rel.get('Type', ''):
            name = rel.get('Target').split('/')[-1]
            return f'xl/drawings/{name}'
    return None


def get_rid_map(zf, drawing_file):
    """drawing.xml.rels から rId → xl/media/... のマップを返す"""
    rels_path = drawing_file.replace('drawings/', 'drawings/_rels/') + '.rels'
    if rels_path not in zf.namelist():
        return {}
    with zf.open(rels_path) as f:
        rels_root = ET.parse(f).getroot()
    return {
        rel.get('Id'): 'xl/media/' + rel.get('Target', '').split('/')[-1]
        for rel in rels_root.findall(q(NS_PKG, 'Relationship'))
        if '../media/' in rel.get('Target', '')
    }


def get_shared_strings(zf):
    """sharedStrings.xml から文字列リストを返す。なければ空リスト"""
    if 'xl/sharedStrings.xml' not in zf.namelist():
        return []
    with zf.open('xl/sharedStrings.xml') as f:
        root = ET.parse(f).getroot()
    shared = []
    for si in root.findall(q(NS_MAIN, 'si')):
        t = si.find(q(NS_MAIN, 't'))
        if t is not None and t.text:
            shared.append(t.text)
        else:
            # リッチテキスト（<r><t>...）の場合は全 run を結合
            shared.append(''.join(
                x.text or '' for x in si.findall(f'.//{q(NS_MAIN, "t")}')
            ))
    return shared


def get_date_fmt_ids(zf):
    """
    styles.xml を解析して「日付フォーマット」と判定される numFmtId の集合を返す。
    Excel 組み込み日付 ID（14-17, 22）＋カスタム numFmt の年月日キーワード検出。
    """
    # Excel 組み込み日付フォーマット ID
    # 14-22: 標準日付, 27-36: 日本語日付, 45-47: 時刻, 50-58: 和暦含む追加日付
    DATE_BUILTIN = set(range(14, 23)) | set(range(27, 37)) | set(range(45, 48)) | set(range(50, 59))
    if 'xl/styles.xml' not in zf.namelist():
        return DATE_BUILTIN, []

    with zf.open('xl/styles.xml') as f:
        root = ET.parse(f).getroot()

    custom_ids = set()
    for nf in root.findall(f'.//{q(NS_MAIN, "numFmt")}'):
        fmt_id   = int(nf.get('numFmtId', 0))
        fmt_code = nf.get('formatCode', '')
        if any(c in fmt_code for c in ['y', 'Y', 'm', 'M', 'd', 'D', '年', '月', '日']):
            custom_ids.add(fmt_id)

    # cellXfs → style index ごとの numFmtId リスト
    xfs = [
        int(xf.get('numFmtId', 0))
        for xf in root.findall(f'.//{q(NS_MAIN, "cellXfs")}/{q(NS_MAIN, "xf")}')
    ]
    return DATE_BUILTIN | custom_ids, xfs


def serial_to_date_str(serial_val):
    """
    Excel 日付シリアル値（数値文字列）を 'YYYY-MM-DD' 形式に変換する。
    変換できない場合は元の文字列をそのまま返す。
    """
    from datetime import date, timedelta
    try:
        n = int(float(serial_val))
        # Excel の 1900 年バグ（60 = 1900-02-29 は存在しない）を考慮
        if n < 1:
            return serial_val
        d = date(1899, 12, 30) + timedelta(days=n)
        return d.strftime('%Y-%m-%d')
    except Exception:
        return serial_val


# ページヘッダー検出パターン
# 文書番号行・ページ番号・作成日をまとめて除外
HEADER_CELL_PATTERN = re.compile(
    r'[Ｎ№N].{0,30}(標準書|作業書|手順書|仕様書)'  # 文書番号
    r'|^[Ｐp]age\s*\d'                               # Page 1/n
    r'|^Ｐａｇｅ'                                    # Ｐａｇｅ（全角）
    r'|^\d{4}年.{1,20}作成$'                         # 2017年 ○月作成
    r'|^２０\d{2}年.{1,20}作成$'                     # ２０１７年...作成
)


def get_cell_anchors(zf, sheet_file, shared_strings):
    """
    sheet XML のセル値を読んでアンカーリストに変換する。

    除外ルール:
      - ページヘッダー行（文書番号・ページ番号・作成日パターン）
      - 同一行に複数セルあり、かつ全セルがヘッダーパターンに合致する行
      - 単独の純粋な数値（小数点のみ可）でテキストとして意味を持たないもの

    変換ルール:
      - 日付フォーマットの数値セル → YYYY-MM-DD 文字列
      - 同一行の複数セル（ヘッダーでない） → kind='table_row'
      - 単独セル → kind='text'（見出しパターン判定あり）
    """
    sheet_path = f'xl/{sheet_file}'
    if sheet_path not in zf.namelist():
        return []

    date_fmt_ids, xfs = get_date_fmt_ids(zf)

    with zf.open(sheet_path) as f:
        root = ET.parse(f).getroot()

    # 行ごとにセル値を収集
    rows = {}  # row_num → [(col_ref, val), ...]
    for row_elem in root.findall(f'.//{q(NS_MAIN, "row")}'):
        row_num = int(row_elem.get('r'))
        for c in row_elem.findall(q(NS_MAIN, 'c')):
            t_attr = c.get('t')
            s_attr = c.get('s')
            v      = c.find(q(NS_MAIN, 'v'))
            is_    = c.find(q(NS_MAIN, 'is'))
            val = None

            if t_attr == 's' and v is not None:
                # 共有文字列
                idx = int(v.text)
                if idx < len(shared_strings):
                    val = shared_strings[idx]
            elif t_attr == 'inlineStr' and is_ is not None:
                val = ''.join(x.text or '' for x in is_.findall(f'.//{q(NS_MAIN, "t")}'))
            elif v is not None and v.text:
                raw = v.text.strip()
                # 日付シリアル値判定
                fmt_id = xfs[int(s_attr)] if s_attr and int(s_attr) < len(xfs) else 0
                if fmt_id in date_fmt_ids and re.fullmatch(r'\d+(\.\d+)?', raw):
                    val = serial_to_date_str(raw)
                else:
                    val = raw

            if val and str(val).strip():
                col_ref = re.sub(r'\d', '', c.get('r', ''))  # 列記号を抽出
                rows.setdefault(row_num, []).append((col_ref, str(val).strip()))

    anchors = []
    for row_num in sorted(rows):
        pairs = rows[row_num]
        vals  = [p[1] for p in pairs]

        # ── ヘッダー行の除外 ────────────────────────────────────────────────
        # 行内の全セルがヘッダーパターンに合致 → スキップ
        if all(HEADER_CELL_PATTERN.search(v) for v in vals):
            continue

        # ── 単独セルで純粋な数値のみ → スキップ ────────────────────────────
        if len(vals) == 1 and re.fullmatch(r'[\d.]+', vals[0]):
            continue

        # 0-indexed row（drawing アンカーとソートキーを合わせる）
        row_idx = row_num - 1

        if len(vals) == 1:
            text = vals[0]
            is_heading = bool(HEADING_PATTERN.match(text.strip()))
            anchors.append({
                'kind': 'text',
                'row': row_idx,
                'paras': [[{'text': text, 'bold': False, 'sz_pt': 10.5, 'color': None}]],
                'is_heading': is_heading,
            })
        else:
            # 複数セル → テーブル行
            anchors.append({
                'kind': 'table_row',
                'row': row_idx,
                'values': vals,
            })

    return anchors


def parse_anchors(zf, drawing_file, rid_map):
    """
    drawing XML を解析してアンカーリストを返す（行番号順にソート済み）。

    各アンカーの構造:
        kind='image' : image_data(bytes), cx_emu(int), cy_emu(int), media_path(str)
        kind='text'  : paras(list), is_heading(bool)
            paras = [ [ {'text', 'bold', 'sz_pt', 'color'} ] ]
    """
    with zf.open(drawing_file) as f:
        root = ET.parse(f).getroot()

    anchors = []

    for anchor in root.findall(q(NS_XDR, 'twoCellAnchor')):
        from_e = anchor.find(q(NS_XDR, 'from'))
        row = int(from_e.find(q(NS_XDR, 'row')).text)

        # ── 画像アンカー ──────────────────────────────────────────────────────
        blip = anchor.find(f'.//{q(NS_A, "blip")}')
        if blip is not None:
            rid = blip.get(q(NS_R, 'embed'))
            media_path = rid_map.get(rid)
            if not media_path or media_path not in zf.namelist():
                continue
            image_data = zf.read(media_path)
            cx_emu = cy_emu = 0
            spPr = anchor.find(f'.//{q(NS_XDR, "spPr")}')
            if spPr is not None:
                xfrm = spPr.find(q(NS_A, 'xfrm'))
                if xfrm is not None:
                    ext = xfrm.find(q(NS_A, 'ext'))
                    if ext is not None:
                        cx_emu = int(ext.get('cx', 0))
                        cy_emu = int(ext.get('cy', 0))
            anchors.append({
                'kind': 'image', 'row': row,
                'image_data': image_data,
                'cx_emu': cx_emu, 'cy_emu': cy_emu,
                'media_path': media_path,
            })
            continue

        # ── テキストアンカー ──────────────────────────────────────────────────
        paras = []
        for p_elem in anchor.findall(f'.//{q(NS_A, "p")}'):
            runs = []
            for r_elem in p_elem.findall(q(NS_A, 'r')):
                t = r_elem.find(q(NS_A, 't'))
                if t is None or not t.text:
                    continue
                rPr = r_elem.find(q(NS_A, 'rPr'))
                bold  = rPr is not None and rPr.get('b') == '1'
                sz_pt = sz_to_pt(rPr.get('sz') if rPr is not None else None)
                color = None
                if rPr is not None:
                    clr = rPr.find(f'.//{q(NS_A, "srgbClr")}')
                    if clr is not None:
                        color = clr.get('val')
                runs.append({'text': t.text, 'bold': bold, 'sz_pt': sz_pt, 'color': color})
            if runs:
                paras.append(runs)

        if paras:
            full_text = ''.join(r['text'] for p in paras for r in p)
            is_heading = bool(HEADING_PATTERN.match(full_text.strip()))
            anchors.append({
                'kind': 'text', 'row': row,
                'paras': paras, 'is_heading': is_heading,
            })

    anchors.sort(key=lambda x: x['row'])
    return anchors


# ─── Word 組み立て ─────────────────────────────────────────────────────────────
def setup_document():
    """A4縦・統一余白・空フッターのドキュメントを返す"""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width        = A4_W
    sec.page_height       = A4_H
    sec.top_margin        = MARGIN_TOP
    sec.bottom_margin     = MARGIN_BOTTOM
    sec.left_margin       = MARGIN_LEFT
    sec.right_margin      = MARGIN_RIGHT
    sec.footer_distance   = Cm(1.0)
    sec.different_first_page_header_footer = False
    # デフォルトの空段落を除去
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return doc


def write_heading(doc, text):
    p = doc.add_heading(text.strip(), level=1)
    for run in p.runs:
        run.font.size      = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)


def write_text(doc, paras):
    for runs in paras:
        if not any(r['text'].strip() for r in runs):
            continue
        p = doc.add_paragraph()
        for ri in runs:
            if not ri['text']:
                continue
            run = p.add_run(ri['text'])
            run.bold       = ri['bold']
            run.font.size  = Pt(ri['sz_pt'])
            if ri['color']:
                try:
                    h = ri['color'].lstrip('#')
                    run.font.color.rgb = RGBColor(
                        int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
                except Exception:
                    pass


def write_image(doc, image_data, cx_emu, cy_emu, media_path):
    """画像をインライン配置。A4本文幅を超える場合はアスペクト比を保って縮小"""
    max_w = CONTENT_W_EMU
    if cx_emu > 0 and cy_emu > 0:
        if cx_emu > max_w:
            cy_emu = int(cy_emu * max_w / cx_emu)
            cx_emu = max_w
        width, height = Emu(cx_emu), Emu(cy_emu)
    else:
        width, height = Emu(max_w), None

    try:
        p = doc.add_paragraph()
        p.add_run().add_picture(io.BytesIO(image_data), width=width, height=height)
    except Exception as e:
        doc.add_paragraph(f'[画像読み込みエラー: {Path(media_path).name} — {e}]')


def write_table_rows(doc, rows_buffer):
    """
    連続する table_row アンカーをまとめて Word テーブルとして出力する。
    rows_buffer: [{'values': [str, ...]}, ...]
    """
    if not rows_buffer:
        return
    col_count = max(len(r['values']) for r in rows_buffer)
    from docx.shared import Pt as _Pt
    from docx.oxml.ns import qn as _qn
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=len(rows_buffer), cols=col_count)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows_buffer):
        for j, val in enumerate(row_data['values']):
            cell = table.cell(i, j)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = _Pt(9)


# ─── シート変換 ────────────────────────────────────────────────────────────────
def convert_sheet(zf, sheet_name, sheet_file, drawing_file, out_dir):
    rid_map      = get_rid_map(zf, drawing_file)
    shared       = get_shared_strings(zf)
    cell_anchors = get_cell_anchors(zf, sheet_file, shared)
    draw_anchors = parse_anchors(zf, drawing_file, rid_map)

    # セルアンカーと drawing アンカーを行番号順にマージ
    anchors = sorted(cell_anchors + draw_anchors, key=lambda x: x['row'])

    if not anchors:
        log(f'スキップ（コンテンツなし）: {sheet_name}', indent=1)
        return False

    doc = setup_document()

    # table_row が連続する場合はまとめてテーブル化するためバッファを使う
    table_buf = []

    def flush_table():
        if table_buf:
            write_table_rows(doc, table_buf)
            table_buf.clear()

    for anchor in anchors:
        if anchor['kind'] == 'table_row':
            table_buf.append(anchor)
        else:
            flush_table()
            if anchor['kind'] == 'image':
                write_image(doc, anchor['image_data'],
                            anchor['cx_emu'], anchor['cy_emu'], anchor['media_path'])
            elif anchor['kind'] == 'text':
                if anchor['is_heading']:
                    full = ''.join(r['text'] for p in anchor['paras'] for r in p)
                    write_heading(doc, full)
                else:
                    write_text(doc, anchor['paras'])

    flush_table()  # 末尾の table_row を処理

    out_path = out_dir / f'{safe_filename(sheet_name)}.docx'
    doc.save(str(out_path))
    log(f'✓ {out_path.name}  ({len(anchors)} アンカー)', indent=2)
    return True


# ─── ファイル変換 ──────────────────────────────────────────────────────────────
def convert_file(xlsx_path, output_root):
    """1つの xlsx を処理。シートごとのサブフォルダに出力"""
    stem = xlsx_path.stem  # 拡張子なしファイル名
    out_dir = output_root / safe_filename(stem)
    out_dir.mkdir(parents=True, exist_ok=True)

    log(f'[{xlsx_path.name}]')

    try:
        with zipfile.ZipFile(xlsx_path, 'r') as zf:
            sheets = get_sheets(zf)
            log(f'シート数: {len(sheets)}', indent=1)
            converted = 0
            for sheet in sheets:
                name         = sheet['name']
                drawing_file = get_drawing_file(zf, sheet['sheet_file'])
                if drawing_file is None or drawing_file not in zf.namelist():
                    log(f'スキップ（drawing なし）: {name}', indent=2)
                    continue
                log(f'処理中: {name}', indent=1)
                if convert_sheet(zf, name, sheet['sheet_file'], drawing_file, out_dir):
                    converted += 1
        log(f'→ {converted} ファイルを output/{safe_filename(stem)}/ に保存', indent=1)
    except Exception as e:
        log(f'エラー: {e}', indent=1)
        raise

    return converted


# ─── エントリーポイント ────────────────────────────────────────────────────────
def main():
    base        = Path(__file__).parent
    input_dir   = base / 'input'
    output_dir  = base / 'output'

    # input/ がなければ作成してガイドを表示
    if not input_dir.exists():
        input_dir.mkdir()
        print('input/ フォルダを作成しました。変換したい .xlsx をここに入れてください。')
        return

    # ~$*.xlsx は Excel が生成するロックファイルのため除外
    xlsx_files = sorted(f for f in input_dir.glob('*.xlsx')
                        if not f.name.startswith('~$'))

    if not xlsx_files:
        print('input/ に .xlsx ファイルが見つかりません。')
        return

    output_dir.mkdir(exist_ok=True)

    print(f'=== xlsx → docx 変換  {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} ===')
    print(f'対象: {len(xlsx_files)} ファイル')
    print()

    total = 0
    for xlsx_path in xlsx_files:
        total += convert_file(xlsx_path, output_dir)
        print()

    print(f'=== 完了: 合計 {total} ファイルを出力 ===')


if __name__ == '__main__':
    main()